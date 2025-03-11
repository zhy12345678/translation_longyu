from flask import Flask, request, jsonify, send_file
import requests
import json
import docx
import re
import os 
from io import BytesIO
from deepdoc.parser.utils import get_text
from rag.nlp import num_tokens_from_string
from deepdoc.parser.txt_parser import RAGFlowTxtParser as TxtParser
from deepdoc.parser.docx_parser import RAGFlowDocxParser as DocxParser
import logging 
import time

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s-%(levelname)s-%(message)s'
)

app = Flask(__name__)

# LOG_FOLDER = os.path.join(os.getcwd(),"logs")
LOG_FOLDER = "/app/logs"
os.makedirs(LOG_FOLDER,exist_ok=True)

def setup_logger(loginName,fileName):
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    log_filename=f"{loginName}_{fileName}_{timestamp}.log"
    log_filepath=os.path.join(LOG_FOLDER,log_filename)

    logger = logging.getLogger(log_filename)
    logger.setLevel(logging.INFO)
    format='%(asctime)s-%(levelname)s-%(message)s'
    formatter = logging.Formatter(format)
    file_handler = logging.FileHandler(log_filepath)
    file_handler.setFormatter(formatter)
    logger.addHandler(file_handler)
    return logger


UPLOAD_FOLDER=os.path.join(os.getcwd(),"translated_files")
os.makedirs(UPLOAD_FOLDER,exist_ok=True)
url = "http://10.27.12.50:1066/v1/chat/completions"

headers = {
    'Content-Type': 'application/json'
}


translation_prompt = """
<instructions>
当您收到一个学术论文或报告的段落时，您的任务是创建一个中英互译(如果输入中文翻译成英文，如果输入英文翻译成中文)的翻译机器人。在进行翻译时，应尽可能进行直译，保持原文的结构和语义。避免使用意译或自由翻译。在翻译过程中，请注意以下几点：
1. 保持原文的句子结构，除非直译会导致语法错误。
2. 将专业术语和特定表达直译，即使在目标语言中可能有更常见的等效表达。
3. 确保翻译后的文本在语法和拼写上是正确的。
4. 输出的翻译文本中不应包含任何XML标签。
5. 并确保不进行重复翻译。
6. 请保持翻译段落的整体性，不要出现一大段英文里面涵盖几个中文字的情况
</instructions>

<examples>
<example>
输入：在过去的几十年里，人工智能（AI）领域取得了显著的进展。
输出：In the past few decades, there has been remarkable progress in the field of artificial intelligence (AI).
</example>

<example>
输入：实验结果表明，新算法在处理大规模数据集时具有优势。
输出：The experimental results show that the new algorithm has advantages in dealing with large-scale data sets.
</example>
</examples>
"""

translation_prompt_multi = """
<instructions>
您的任务是创建语言翻译机器人。根据输入的语言（`translate_language`）和目标语言（`translated_language`）进行翻译。在翻译过程中，请遵循以下规则：
1. 请自动检测输入文本的语言
2. 需要保持原文的句子结构，除非直译会导致语法错误。将专业术语和特定表达直译，即使在目标语言中可能有更常见的等效表达。
3. 请根据 `translate_language` 和 `translated_language` 确定翻译方向，例如：
   - 输入语言`translate_language`为 `中文` 且目标语言`translated_language`为 `英文`，将中文翻译为英文。
   - 输入语言`translate_language`为 `英文` 且目标语言`translated_language`为 `中文`，将英文翻译为中文。
4. 确保翻译后的文本在语法和拼写上是正确的。
5. 输出的翻译文本中不应包含任何XML标签。
6. 并确保不进行重复翻译。如果输入文本为空则直接返回空字符串
7. 如果遇到检测的输入文本语言和目标语言相通的情况直接返回原文
</instructions>

<examples>
<example>
输入语言：中文
目标语言：英语
输入：在过去的几十年里，人工智能（AI）领域取得了显著的进展。
输出：In the past few decades, there has been remarkable progress in the field of artificial intelligence (AI).
</example>

<example>
输入语言：英语
目标语言：中文
输入： The experimental results show that the new algorithm has advantages in dealing with large-scale data sets.
输出： 实验结果表明，新算法在处理大规模数据集时具有优势。
</example>
</examples>
"""

@app.route('/translate', methods=['POST'])
def translate():
    try:
        data = request.get_json()
        user_input = data.get('text', '')
        translate_language = data.get('translate_language', '')  # 输入语言
        translated_language = data.get('translated_language', '')  # 目标语言
        if not translate_language or not translated_language:
            return jsonify({"error": "Both translate_language and translated_language are required."}), 400
        
        prompt = f"""{translation_prompt_multi}
<example>
输入语言：{translate_language}
目标语言：{translated_language}
输入文本：{user_input}
输出（仅提供翻译，无需其他内容）：
</example>
"""

        payload = json.dumps({
            "model": "/root/models/glm4/qwen2.5-32b-int4/", 
            "stream": False,
            "messages": [
                {
                    "role": "user",
                    "content": prompt
                }
            ]
        })

        response = requests.post(url, headers=headers, data=payload)
        result = response.json()

        translation = result['choices'][0]['message']['content']
        split_translation = translation.split("\n\n")
        translation = "\n\n".join(split_translation)
        if len(split_translation) > 1:
            translation = split_translation[0]
        return jsonify({"translation": translation})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/translate_text', methods=['POST'])
def translate_text():
    try:
        data = request.get_json()
        user_input = data.get('text', '')
        loginName=data.get('loginName')
        translate_language = data.get('translate_language')  
        translated_language = data.get('translated_language') 
        fileName="text"
        
        logger=setup_logger(loginName,fileName)
        logger.info(f"Received file from user {loginName}")

        if not user_input.strip():
            logger.error("Input text is empty.")
            return jsonify({"error": "Input text is empty.", "error_code": "801"}), 400
        try:
            logger.info("Sending full text to translation API.")
            prompt = f"""{translation_prompt_multi}
<example>
输入语言：{translate_language}
目标语言：{translated_language}
输入文本：{user_input}
输出（仅提供翻译，无需其他内容）：
</example>
"""
            payload = json.dumps({
            "model": "/root/models/glm4/qwen2.5-32b-int4/", 
            "stream": False,
            "messages": [
                {
                    "role": "user",
                    "content": prompt
                }
            ]
        })
            response = requests.post(url, headers=headers, data=payload)
            result = response.json()

            translation = result['choices'][0]['message']['content']
            return jsonify({"translation": translation})

        except requests.RequestException as req_err:
            logger.error(f"Request error: {str(req_err)}")
            return jsonify({"error": "Translation request failed.", "details": str(req_err)}), 500
        except Exception as e:
            logger.error(f"Unexpected error during translation: {str(e)}")
            return jsonify({"error": "Internal server error.", "details": str(e)}), 500

    except Exception as e:
        logging.error(f"Unexpected error: {str(e)}")
        return jsonify({"error": "Internal server error.", "details": str(e)}), 500

# @app.route('/translate_text', methods=['POST'])
# def translate_text():
#     try:
#         data = request.get_json()
#         user_input = data.get('text', '')
#         loginName=data.get('loginName')
#         translate_language = data.get('translate_language')  
#         translated_language = data.get('translated_language') 
#         fileName="text"
        
#         logger=setup_logger(loginName,fileName)
#         logger.info(f"Received file from user {loginName}")

#         # total_character_count = len(file_content)
#         # if total_character_count > 50000:
#         #     logger.error(f"File contains too many characters: {total_character_count}")
#         #     return jsonify({
#         #         "error": "File contains more than 50,000 characters.",
#         #         "error_code": "800",  
#         #         "character_count": total_character_count  
#         #         }), 400

#         chunks = user_input.split("\n")
            
#         translated_chunks = []  
#         for chunk in chunks:
#             text_to_translate = chunk.strip()
#             logger.info(f"start translate:{text_to_translate[0:15]}...")
#             if text_to_translate == "":
#                 translated_chunks.append("")
#                 continue
#             try:
#                 payload = {
#                     "text": text_to_translate,
#                      "translate_language": translate_language,
#                      "translated_language": translated_language
#                      }
#                 response = requests.post("http://127.0.0.1:5000/translate", json=payload)
#                 result = response.json()
#                 if "translation" in result:
#                     translationed=result["translation"]
#                     translated_chunks.append(translationed)  
#                 else:
#                     logger.info(f"翻译接口返回错误信息")
#             except Exception as e:
#                 translated_chunks.append(f"Error translating chunk: {str(e)}")  
#                 logger.info(f"Error translating chunk: {str(e)}")

#         final_translation = "\n".join(translated_chunks)  
#         logger.info(f"Translation complete successful")   
#         return jsonify({"translation": final_translation}) 
        
#     except Exception as e:
#         return jsonify({"error": str(e)}), 500  


@app.route('/translate_txt', methods=['POST'])
def translate_txt():
    try:
        file = request.files.get('document')
        loginName=request.form.get('loginName')
        translate_language = request.form.get('translate_language')  
        translated_language = request.form.get('translated_language') 
        
        if not file:
            return jsonify({"error": "No file uploaded"}), 400
        file_name = file.filename
        logger=setup_logger(loginName,file_name)
        logger.info(f"Received file:{file_name} from user {loginName}")
        file_content=file.read().decode('utf-8')

        chunks = file_content.split("\n")
            
        translated_chunks = []  
        for chunk in chunks:
            text_to_translate = chunk.strip()
            logger.info(f"start translate:{text_to_translate[0:5]}...")
            if text_to_translate == "":
                translated_chunks.append("")
                continue
            try:
                payload = {
                    "text": text_to_translate,
                     "translate_language": translate_language,
                     "translated_language": translated_language
                     }
                response = requests.post("http://127.0.0.1:5000/translate", json=payload)
                result = response.json()
                if "translation" in result:
                    translationed=result["translation"]
                    combined_text=f"{text_to_translate}\n{translationed}"
                    translated_chunks.append(combined_text)  
                else:
                    logger.info(f"翻译接口返回错误信息")
            except Exception as e:
                translated_chunks.append(f"Error translating chunk: {str(e)}")  
                logger.info(f"Error translating chunk: {str(e)}")

        final_translation = "\n".join(translated_chunks)  
        logger.info(f"Translation complete successful")   
        return jsonify({"translation": final_translation}) 
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500  




@app.route('/translate_docx_plain', methods=['POST'])
def translate_docx_plain():
    try:
        file = request.files.get('document')
        logging.info(f"translation start")
        original_file_name = file.filename.rsplit('.', 1)[0]
        logging.info(f"translation get file name ")
        if not file:
            return jsonify({"error": "No file uploaded"}), 400
        
        parser = DocxParser()
        file_content = file.read()
        sections, tables = parser(file_content)  
        
        translated_doc = docx.Document()  

        for section in sections:
            translated_text = ""
            try:
                if section[0].strip(): 
                    text_to_translate = section[0]
                    logging.info(f"text_to_translate{text_to_translate}")
                    payload = {"text": text_to_translate}
                    response = requests.post("http://127.0.0.1:5000/translate", json=payload)
                    result = response.json()
                    if "translation" in result:
                        logging.info(f"translation{result['translation']}")
                        translated_text += result["translation"]  
                    else:
                        print(f"翻译接口返回错误信息: {result}")

                    para = translated_doc.add_paragraph(text_to_translate)
                    para.style = section[1]  
                    para = translated_doc.add_paragraph(translated_text)
                    para.style = section[1]  
                else:
                    translated_text = ""
                    # translated_text = section[0]  
                
            except Exception as e:
                translated_text = ""
                para = translated_doc.add_paragraph(translated_text)
                para.style = 'Normal'

        
        for table in tables: 
            num_cols=len(table[0].split(";")) if len(table)>0 else 0
            tbl = translated_doc.add_table(rows=0, cols=num_cols)  
            for row in table:
                tbl_row = tbl.add_row()
                cells = row.split(";")  
                for idx, cell in enumerate(cells):
                    text_to_translate=cell.strip()
                    if text_to_translate:
                        payload = {"text" : text_to_translate}
                        response = requests.post("http://127.0.0.1:5000/translate", json=payload)
                        result = response.json()
                        if "translation" in result:
                            translated_text = result["translation"]
                            tbl_row.cells[idx].text = f"Original:{text_to_translate}\nTranslated:{translated_text}"
                            logging.info(f"translation{result['translation']}")
                        else:
                            tbl_row.cells[idx].text = f"Error:{result}"
                    else:
                        tbl_row.cells[idx].text = "Empty"

        output_stream = BytesIO()
        translate_file_path=os.path.join(UPLOAD_FOLDER,f'{original_file_name}_translate.docx')
        logging.info(f"translate_file_path{translate_file_path}")
        translated_doc.save(translate_file_path)
        translated_doc.save(output_stream)
        output_stream.seek(0)

        return send_file(output_stream, as_attachment=True, download_name=translate_file_path, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500



@app.route('/translate_docx', methods=['POST'])
def translate_docx():
    try:
        file = request.files.get('document')
        loginName=request.form.get('loginName')
        translate_language = request.form.get('translate_language')  
        translated_language = request.form.get('translated_language')  

        file_name = file.filename
        original_file_name = file.filename.rsplit('.', 1)[0]
        logger=setup_logger(loginName,file_name)
        logger.info(f"Received file:{file_name} from user {loginName}")
        logger.info(f"Translation from {translate_language} to {translated_language}")
        if not file:
            return jsonify({"error": "No file uploaded"}), 400
        
        parser = DocxParser()
        file_content = file.read()
        original_doc = docx.Document(BytesIO(file_content))
        logger.info(f"Successfully read DOCX content")
        for paragraph in original_doc.paragraphs:
            if paragraph.text.strip():  
                text_to_translate = paragraph.text.strip()
                logger.info(f"start translate:{text_to_translate[0:10]}...")
                payload = {
                    "text": text_to_translate,
                    "translate_language": translate_language,
                    "translated_language": translated_language
                    }
                try:
                    response = requests.post("http://127.0.0.1:5000/translate", json=payload)
                    result = response.json()
                    if "translation" in result:
                        translated_text = result["translation"]
                        paragraph.add_run("\n")
                        paragraph.add_run(translated_text)
                        logger.info(f"translated content:{translated_text[0:10]}...")
                    else:
                        logger.warning(f"translation API returned an error:{result}")
                except Exception as e:
                    logger.error(f"Error during translation:{str(e)}")
                    paragraph.add_run("\nError: Translation failed.")
        
        for table in original_doc.tables: 
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_to_translate=cell.text.strip()
                        logger.info(f"start translate...")
                        payload = {
                            "text" : text_to_translate,
                            "translate_language": translate_language,
                            "translated_language": translated_language
                            }
                        try:
                            response = requests.post("http://127.0.0.1:5000/translate", json=payload)
                            result = response.json()
                            if "translation" in result:
                                translated_text = result["translation"]
                                cell.text = f"{text_to_translate}\n{translated_text}"
                                logger.info(f"translated content...")
                            else:
                                logger.warning(f"translation API returned an error")
                        except Exception as e:
                            logger.error(f"error during translation{str(e)}")
                            cell.text += "\nError"
               

        output_stream = BytesIO()
        translate_file_path=os.path.join(UPLOAD_FOLDER,f'{original_file_name}_translate.docx')
        logger.info(f"translate file path:{translate_file_path}")
        original_doc.save(translate_file_path)
        original_doc.save(output_stream)
        output_stream.seek(0)
        logger.info(f"Translation complete successful")  

        return send_file(output_stream, as_attachment=True, download_name=translate_file_path, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    
    except Exception as e:
        logger.error(f"Unexpected error{str(e)}")
        return jsonify({"error": str(e)}), 500



@app.route('/translate_document', methods=['POST'])
def translate_document():
    try:
        file = request.files.get('document')
        if file and file.filename.endswith('.docx'):
            original_file_name = file.filename.rsplit('.', 1)[0]
            doc = docx.Document(file)
            translated_text_chunks = []
            paragraphs_buffer = [] 
            is_heading_section = False  
            for index, p in enumerate(doc.paragraphs):
                style_name = p.style.name
                paragraphs_buffer.append(p)
                if style_name.startswith('Heading'):
                    is_heading_section = True
                    continue
                else:
                    if len(paragraphs_buffer) >= 2 or (is_heading_section and index == len(doc.paragraphs) - 1):
                        text_to_translate = ""
                        for para in paragraphs_buffer:
                            text_to_translate += para.text + "\n"

                        payload = {
                            "text": text_to_translate
                        }
                        response = requests.post("http://127.0.0.1:5000/translate", json=payload)
                        result = response.json()
                        if "translation" in result:
                            translated_text_chunks.append(result["translation"])
                        else:
                            print(f"翻译接口返回错误信息: {result}")
                        paragraphs_buffer = []  
                        is_heading_section = False
                    

            translated_doc = docx.Document()
            for translated_chunk in translated_text_chunks:
                translated_doc.add_paragraph(translated_chunk)

            translate_file_path=os.path.join(UPLOAD_FOLDER,f'{original_file_name}_translate.docx')
            print(translate_file_path)
            translated_doc.save(translate_file_path)
            file_url=f'/download/{original_file_name}_translate.docx'
           
            return jsonify({"message": "文档翻译并保存成功","file_url":file_url})
        else:
            return jsonify({"error": "请上传有效的Word文档"}), 400
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/download/<path:filename>',methods=['GET'])
def download(filename):
    try:
        return send_from_directory(UPLOAD_FOLDER,filename,as_attachment=True)
    except Exception as e:
        return jsonify({"error": str(e)}),500


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
